# %% [markdown]
# Bibliotecas

# %% [markdown]
# # Nova seção

# %% [markdown]
# Instalando Bibliotecas

# %%
!pip install python-docx
!pip install docx
!pip install pandas
!pip install beautifulsoup4
!pip install ipywidgets
!pip install openpyxl


# %%
import requests
import zipfile
import os
import pandas as pd
from bs4 import BeautifulSoup
import ipywidgets as widgets
from IPython.display import display
from docx import Document
from tqdm import tqdm
from datetime import datetime


# %% [markdown]
# Diretórios 

# %%
dir_arq_cvm_zip = r'C:\Users\User\Documents\GitHub\TCC---MBA_USP_-DSA\Arquivos TCC\arquivos_cvm_zip\arquivos_cvm_zip'
lista_empresas_b3_input = r'C:\Users\User\Documents\GitHub\TCC---MBA_USP_-DSA\Arquivos TCC\arquivos_cvm_zip\arquivos_cvm_zip\Outros arquivos\elemento_html_b3.docx'
lista_empresas_b3 = r'C:\Users\User\Documents\GitHub\TCC---MBA_USP_-DSA\Arquivos TCC\arquivos_cvm_zip\arquivos_cvm_zip\Outros arquivos\empresas_listadas_b3.csv'
dir_empresas_disp_cvm = r'C:\Users\User\Documents\GitHub\TCC---MBA_USP_-DSA\Arquivos TCC\arquivos_cvm_zip\arquivos_cvm_zip\Outros arquivos\Empresas_dispensadas_CVM'
dir_financeiras = r'C:\Arquivos TCC\planilhas_excel\planilhas_excel\Financeiras'
dir_out_n_acoes = r'C:\Users\User\Documents\GitHub\TCC---MBA_USP_-DSA\Arquivos TCC\empresas_cvm_tcc_csv-'

# %% [markdown]
# ## **Extrair os arquivos da CVM**

# %%
# URL base da API
url_base = 'http://dados.cvm.gov.br/dados/CIA_ABERTA/DOC/DFP/DADOS'

# Diretório de destino para salvar os arquivos ZIP
diretorio_destino = dir_arq_cvm_zip
# Verificando se o diretório de destino existe
os.makedirs(diretorio_destino, exist_ok=True)

# Anos em que os relatórios serão obtidos
ano_inicial =2012
ano_final=2023

for ano in range(ano_inicial, ano_final + 1):
    relatorio = "dfp_cia_aberta"
    arquivo_zip = f'{relatorio}_{ano}.zip'
    caminho_arquivo_local = os.path.join(diretorio_destino, arquivo_zip)

    # Verificando se o arquivo já existe no diretório de destino
    if os.path.exists(caminho_arquivo_local):
        print(f"Arquivo {arquivo_zip} já existe. Pulando o download e extração.")
    else:
        url_arquivo = f'{url_base}/{arquivo_zip}'

        # Download do arquivo ZIP
        response = requests.get(url_arquivo)
        if response.status_code == 200:
            with open(caminho_arquivo_local, 'wb') as arquivo_local:
                arquivo_local.write(response.content)

            # Extração do conteúdo do arquivo ZIP
            with zipfile.ZipFile(caminho_arquivo_local, 'r') as zip_ref:
                # Coletar os nomes dos arquivos extraídos
                lista_arquivos_zip = zip_ref.namelist()

                # Exibindo a lista de nomes dos arquivos extraídos
                print("Arquivos no arquivo ZIP:")
                for arquivo in lista_arquivos_zip:
                    print(arquivo)

                # Determinando qual arquivo do ZIP será lido.
                primeiro_arquivo = lista_arquivos_zip[0]
                with zip_ref.open(primeiro_arquivo) as arquivo:
                    df = pd.read_csv(arquivo, encoding='cp1252', sep=';')

                # Exibe as primeiras linhas do DataFrame
                df.head()
        else:
            print(f"Falha ao baixar {arquivo_zip}. Código de status:", response.status_code)



# %% [markdown]
# Visualizando os arquivos obtidos

# %%
# Lista com todos os arquivos no diretório disponíveis para visualização
arquivos_disponiveis = os.listdir(diretorio_destino)

# Exibe a lista de arquivos disponíveis
for arquivo in arquivos_disponiveis:
    print(arquivo)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind" #inserir apenas o nome do arquivo, sem o ano
ano = 2012

# Definindo o caminho do arquivo usando formatação de string
caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')




# %%
lista_ds_conta = df[['CD_CONTA', 'DS_CONTA']].drop_duplicates()

lista_ds_conta = lista_ds_conta.sort_values(by=['DS_CONTA'])

df_lista = pd.DataFrame(lista_ds_conta)

df_lista


# %%
dt = DataTable(df)
dt

# %% [markdown]
# ===============================================================================

# %% [markdown]
# ## **Obtenção das empresas listadas na url da B3**
# 
# 

# %% [markdown]
# 1.1- Extraindo a lista das empresas listadas na B3 pela url: https://sistemaswebb3-listados.b3.com.br/shareCapitalPage/?language=pt-br (acesso em 17/10/2023)

# %%
# URL da página
url = "https://sistemaswebb3-listados.b3.com.br/shareCapitalPage/?language=pt-br"

# Solicitando ao HTTP para obter o conteúdo da página
response = requests.get(url)
content = response.text

# Utilizando o BeautifulSoup para analisar o conteúdo da página
soup = BeautifulSoup(content, "html.parser")

# Encontrando a tabela com os atributos especificados
tabela = soup.find("table", {"cdk-table": "", "class": "cdk-table table table-striped table-hover visible", "id": "shareCapitalTable"})

if tabela:
    print(tabela)
else:
    print("Nenhuma tabela encontrada com os atributos especificados.")


# %% [markdown]
# 1.2- As tabelas não foram encontradas por web scraping, desta forma, os elementos foram obtidos e salvos em uma lista para a obtenção da lista das empresas da B3

# %%
#extração da tabela pelo elemento da tabela disponível no site: https://sistemaswebb3-listados.b3.com.br/shareCapitalPage/?language=pt-br (a tabela não estava sendo identificada pelo método soup.find)

# Abrindo o arquivo .docx
doc = Document(lista_empresas_b3_input) 

# Inicializando uma string vazia para armazenar o conteúdo do arquivo .docx
elemento_html = ""

# Extraindo o texto do documento .docx
for paragraph in doc.paragraphs:
    elemento_html += paragraph.text

# Utilizando o BeautifulSoup para analisar o elemento HTML
soup = BeautifulSoup(elemento_html, "html.parser")

# Encontrando a tabela dentro do elemento
tabela = soup.find("table")

# Usando a função read_html do pandas para converter a tabela em um DataFrame
empresas_listadas_b3 = pd.read_html(str(tabela))[0]


caminho_arquivo = lista_empresas_b3

# Salvando o DataFrame como um arquivo CSV no caminho especificado
empresas_listadas_b3.to_csv(caminho_arquivo, index=False)

# %%
empresas_listadas_b3

# %% [markdown]
# 1.3- Na url da B3 há a separação das empresas que são dispensadas de registro na CVM, contudo, mesmo que não tendo registro na CVM  podem estar sujeitas a certas obrigações de divulgação de informações à CVM.Por este motivos, estas empresas foram incorporadas à listagem para verificação de existencia de relatório.

# %%
# Caminho para os arquivos DOCX
caminho_arquivos = dir_empresas_disp_cvm 

# Lista para armazenar os DataFrames de cada arquivo
dfs = []

# Iterar sobre os arquivos DOCX
for i in range(1, 16):
    # Caminho completo do arquivo
    caminho_arquivo = os.path.join(caminho_arquivos, f'Empresas_dispensadas_CVM_{i}.docx')

    # Abrir o arquivo DOCX
    doc = Document(caminho_arquivo)

    # Extrair o conteúdo HTML do arquivo DOCX
    conteudo_html = ""
    for paragrafo in doc.paragraphs:
        conteudo_html += paragrafo.text

    # Converter o conteúdo HTML para DataFrame
    df = pd.read_html(conteudo_html, header=0)[0]

    # Adicionar o DataFrame à lista
    dfs.append(df)

# Concatenar todos os DataFrames em um único DataFrame
empresas_disp_cvm = pd.concat(dfs, ignore_index=True)

# Verificar o DataFrame consolidado
print(empresas_disp_cvm.head())


# %% [markdown]
# 2.1- Preparando os dfs para jução das listagens das ações listadas na B3

# %%
# Dicionário de mapeamento dos novos nomes de colunas
novo_nome_colunas = {
   'Razão Social': 'Denominação Social',
    'Nome de Pregão':'Nome Do Pregão' ,
    'Segmento':'Segmento De Mercado',
    'Código_b3': 'Código'
}

# Renomear as colunas
empresas_disp_cvm = empresas_disp_cvm.rename(columns=novo_nome_colunas)


# %%
# Extrair as colunas desejadas de "empresas_listadas_b3"
colunas_desejadas = ["Denominação Social", "Nome Do Pregão", "Segmento De Mercado", "Código"]
empresas_listadas_subset = empresas_listadas_b3[colunas_desejadas]

# Concatenar os DataFrames
empresas_totais_b3 = pd.concat([empresas_listadas_subset, empresas_disp_cvm], ignore_index=True)

# Remover as linhas duplicadas com base em todas as colunas
# Concatenar os DataFrames
empresas_totais_b3 = empresas_totais_b3.drop_duplicates()




# %%
empresas_totais_b3

# %% [markdown]
# Neste código abaixo será criado um df com as empresas financeiras listadas na B3. Este será necessário para a esclusão destas empresas do dataset final

# %%
# Diretório onde os arquivos .docx estão localizados
diretorio_destino = dir_financeiras
# Inicializando um DataFrame vazio para armazenar os dados das tabelas
df_final = pd.DataFrame()

# Iterando sobre os arquivos no diretório
for arquivo in os.listdir(diretorio_destino):
    if arquivo.endswith(".docx"):
        caminho_arquivo = os.path.join(diretorio_destino, arquivo)

        # Abrindo o arquivo DOCX
        doc = Document(caminho_arquivo)

        # Inicializando uma string vazia para armazenar o conteúdo HTML
        elemento_html = ""

        # Extraindo o texto do documento .docx
        for paragraph in doc.paragraphs:
            elemento_html += paragraph.text

        # Utilizando o BeautifulSoup para analisar o elemento HTML
        soup = BeautifulSoup(elemento_html, "html.parser")

        # Encontrando a tabela dentro do elemento
        tabela = soup.find("table")

        # Verificando se a tabela foi encontrada
        if tabela:
            # Usar a função read_html do Pandas para converter a tabela em um DataFrame
            df = pd.read_html(str(tabela))[0]

            # Adicionar os dados do DataFrame à DataFrame final
            df_financeiras = pd.concat([df_final, df], ignore_index=True)
# Adicionar os dados do DataFrame à DataFrame final
            df_final = pd.concat([df_final, df], ignore_index=True)
            df_financeiras = df_final

df_financeiras.to_csv(os.path.join(dir_financeiras, 'financeiras.csv'), index=False)




# %%
nome_arquivo_excel = 'financeiras.xlsx'
nome_folha = 'Página1'
caminho_arquivo_excel = os.path.join(dir_financeiras, nome_arquivo_excel)
#nomes_planilhas = pd.ExcelFile(caminho_arquivo_excel).sheet_names
#print("Nomes das Planilhas:", nomes_planilhas)
df_financeiras = pd.read_excel(caminho_arquivo_excel, sheet_name=nome_folha)

# %%
df_financeiras

# %% [markdown]
# Removendo a lista de empresas financeiras do df empresas_totais_b3

# %%
# Coluna 'Código' do DataFrame df_financeiras
coluna_codigo_financeiras = 'Código'

# Lista de códigos para remover
codigos_a_remover = df_financeiras[coluna_codigo_financeiras].tolist()

# Filtrar o DataFrame empresas_totais_b3
empresas_totais_b3 = empresas_totais_b3[~empresas_totais_b3['Código'].isin(codigos_a_remover)]


# %%
empresas_totais_b3

# %% [markdown]
# A partir  da listagem das empresas da B3, será realizada a comparação com o df das empresas que possuem relatórios na CVM para a obtenção dos CNPJ e posterior obtenção dos dados

# %% [markdown]
# 2.1- Criar uma lista com os nomes das empresas com relatórios na CVM

# %%
# Lista para armazenar os nomes das empresas
nomes_empresas = []

# Coletando os nomes das empresas no DataFrame específico
for arquivo in os.listdir(diretorio_destino):
    if arquivo.endswith(".csv"):
        caminho_arquivo = os.path.join(diretorio_destino, arquivo)
        if 'dfp_cia_aberta_DRE_ind' in caminho_arquivo:
            df = pd.read_csv(caminho_arquivo, sep=';', encoding='latin1')
            # Verificando se o DataFrame contém a coluna 'DENOM_CIA'
            if 'DENOM_CIA' in df.columns:
                nomes_empresas.extend(df['DENOM_CIA'].unique())

# Removendo duplicatas
nomes_empresas = list(set(nomes_empresas))


# %%
#abrindo um relatório aleatório para buscar os cnpjs
# Definindo o caminho do arquivo usando formatação de string
documento = 'dfp_cia_aberta_DRE_con_2022.csv'
caminho_arquivo = os.path.join(dir_arq_cvm_zip, documento)
dra_con_2012 = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')


#Obtendo a Denominação Social da empresas da B3
razao_social = empresas_totais_b3.sort_values(by='Denominação Social', ascending=True)

# Criando o DataFrame `cnpj_empresas_b3` vazio
cnpj_empresas_b3 = pd.DataFrame(columns=['Denominação Social', 'CNPJ'])
doc_out = 'cnpj_empresas_b3.csv'
caminho_saida = os.path.join(dir_arq_cvm_zip, doc_out)

# Carregando o arquivo CSV se existir, caso contrário, criará um DataFrame vazio
try:
    cnpj_empresas_b3 = pd.read_csv(caminho_saida)
except FileNotFoundError:
    cnpj_empresas_b3 = pd.DataFrame(columns=['Denominação Social', 'CNPJ'])

# Iterar pelas linhas do DataFrame `empresas_totais_b3`
for index, row in empresas_totais_b3.iterrows():
    # Verificar se a 'Denominação Social' já existe no DataFrame `cnpj_empresas_b3`
    existe = cnpj_empresas_b3['Denominação Social'] == row['Denominação Social']

    # Se a 'Denominação Social' existir, atualize o 'CNPJ' correspondente
    if existe.any():
        indice = existe[existe].index[0]
        cnpj_empresas_b3.at[indice, 'CNPJ'] = dra_con_2012.loc[dra_con_2012['DENOM_CIA'] == row['Denominação Social'], 'CNPJ_CIA'].iloc[0]
    else:
        # Caso contrário, adicione um novo registro ao DataFrame `cnpj_empresas_b3`
        cnpj = dra_con_2012.loc[dra_con_2012['DENOM_CIA'] == row['Denominação Social'], 'CNPJ_CIA']
        if not cnpj.empty:
            cnpj_empresas_b3 = cnpj_empresas_b3.append({'Denominação Social': row['Denominação Social'], 'CNPJ': cnpj.iloc[0]}, ignore_index=True)

# Salvando o DataFrame atualizado em um arquivo CSV
cnpj_empresas_b3.to_csv(caminho_saida, index=False)

# %%
cnpj_empresas_b3

# %% [markdown]
# ######Este código foi criado com o intuito de vereificar se alguma Denominação Social apresentava diferença nos arquivos da CVM. Ele faz um split no mome e compara se ha´ao menos 2 palavras no nome que são iguais. Mas o tempo de execução, em função das iterações linha a linha, é alto (foi de 31min).

# %%
'''# Crie o DataFrame `cnpj_empresas_b3` vazio
cnpj_empresas_b3_2 = pd.DataFrame(columns=['Denominação Social', 'CNPJ'])

# Itere pelas linhas do DataFrame `empresas_totais_b3`
for index, row in empresas_totais_b3.iterrows():
    # Divida a "Denominação Social" em palavras
    palavras_empresas = set(row['Denominação Social'].split())

    # Itere pelo DataFrame `dra_con_2012` e encontre o CNPJ correspondente com base nas palavras iguais
    for _, row_dra_con in dra_con_2012.iterrows():
        palavras_dra_con = set(row_dra_con['DENOM_CIA'].split())

        # Verifique se há pelo menos duas palavras iguais entre as duas strings
        if len(palavras_empresas & palavras_dra_con) >= 2:
            # Adicione o CNPJ correspondente ao DataFrame `cnpj_empresas_b3`
            cnpj_empresas_b3_2 = cnpj_empresas_b3.append({'Denominação Social': row['Denominação Social'], 'CNPJ': row_dra_con['CNPJ_CIA']}, ignore_index=True)'''



# %%
#cnpj_empresas_b3_2

# %% [markdown]
# ================================================================================

# %% [markdown]
# ## **Obtenção das variáveis dentros dos arquivos**

# %% [markdown]
# **<Lucro/ação>** = DRE
# **<Patrimônio líquido>** = BPP_con
# composição da dívida total (passivos onerosos) em relação ao capital total [DTPL] = endividamento (FRE_endividamento)/patrimonio Liquido * 100 e [DTCT] = endividamento (FRE_endividamento)/capital total * 100 .
# retorno sobre o patrimônio líquido (“Return on Equity” [ROE]) = **Lucro Líquido** (DRA) / **Patrimônio líquido**
# <ROA = (**lucro líquido** (DRA) / **ativo total** (BPA) x 100>
# 

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind" #inserir apenas o nome do arquivo, sem o ano
ano = 2012

# Definindo o caminho do arquivo usando formatação de string
caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')




# %%
lista_ds_conta = df[['CD_CONTA', 'DS_CONTA']].drop_duplicates()

lista_ds_conta = lista_ds_conta.sort_values(by=['DS_CONTA'])

df_lista = pd.DataFrame(lista_ds_conta)

df_lista


# %% [markdown]
# Lucro Líquido consolidado do período (DRE)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dre = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_DRE = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '3.11')]

    # Selecionando colunas desejadas
    var_DRE = var_DRE[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DRE['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dre = pd.concat([resultados_dre, var_DRE], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dre = resultados_dre.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
lucro_liquido = resultados_agregados_dre.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
lucro_liquido.columns.name = None


lucro_liquido





# %% [markdown]
# Ativo total (BPA)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPA_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpa = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPA = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '1')]

    # Selecionando colunas desejadas
    var_BPA = var_BPA[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPA['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpa = pd.concat([resultados_bpa, var_BPA], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpa = resultados_bpa.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
ativo_total = resultados_agregados_bpa.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
ativo_total.name = None


ativo_total

# %% [markdown]
# Ativo Circulante

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPA_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpa = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPA = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '1.01')]

    # Selecionando colunas desejadas
    var_BPA = var_BPA[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPA['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpa = pd.concat([resultados_bpa, var_BPA], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpa = resultados_bpa.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
ativo_circulante = resultados_agregados_bpa.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
ativo_circulante.name = None


ativo_circulante

# %% [markdown]
# Patrimônio Líquido (BPP)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPP_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpp = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPP = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '2.03')]

    # Selecionando colunas desejadas
    var_BPP = var_BPP[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPP['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpp = pd.concat([resultados_bpp, var_BPP], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpp = resultados_bpp.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
patrimonio_liquido = resultados_agregados_bpp.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
patrimonio_liquido.name = None


patrimonio_liquido

# %% [markdown]
# Passivo Circulante

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPP_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpp = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPP = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '2.01')]

    # Selecionando colunas desejadas
    var_BPP = var_BPP[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPP['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpp = pd.concat([resultados_bpp, var_BPP], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpp = resultados_bpp.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
passivo_circulante = resultados_agregados_bpp.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
passivo_circulante.name = None


passivo_circulante

# %% [markdown]
# Passivo Exigível a longo Prazo (Passivo não Circulante)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPP_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpp = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')
    
    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPP = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '2.02')]

    # Selecionando colunas desejadas
    var_BPP = var_BPP[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPP['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpp = pd.concat([resultados_bpp, var_BPP], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpp = resultados_bpp.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
passivo_exig_longo_prazo = resultados_agregados_bpp.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
passivo_exig_longo_prazo.name = None


passivo_exig_longo_prazo

# %% [markdown]
# Lucro por ação (DRE)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dre = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_DRE = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '3.99')]

    # Selecionando colunas desejadas
    var_DRE = var_DRE[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DRE['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dre = pd.concat([resultados_dre, var_DRE], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dre = resultados_dre.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
lucro_acao = resultados_agregados_dre.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
lucro_acao.name = None


lucro_acao

# %% [markdown]
# Vendas Líquidas (Receita Líquida)

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dre = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_DRE = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '3.01')]

    # Selecionando colunas desejadas
    var_DRE = var_DRE[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DRE['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dre = pd.concat([resultados_dre, var_DRE], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dre = resultados_dre.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
vendas_liquidas = resultados_agregados_dre.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
vendas_liquidas.name = None

# Selecionando as colunas desejadas
vendas_liquidas = vendas_liquidas[['DENOM_CIA', 'CNPJ_CIA'] + list(range(2012, 2023))]


# %% [markdown]
# EBITIDA

# %% [markdown]
# Depreciação e amortização

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DVA_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dva = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_DVA = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '7.04.01')]

    # Selecionando colunas desejadas
    var_DVA = var_DVA[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DVA['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dva = pd.concat([resultados_dva, var_DVA], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dva = resultados_dva.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
deprec_amortz = resultados_agregados_dva.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
deprec_amortz.name = None

# %% [markdown]
# LAIR

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dre = pd.DataFrame()

# Loop pelos anos
for ano in anos:
   # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_DRE = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '3.05')]

    # Selecionando colunas desejadas
    var_DRE = var_DRE[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DRE['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dre = pd.concat([resultados_dre, var_DRE], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dre = resultados_dre.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
LAIR = resultados_agregados_dre.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
LAIR.name = None



# %%
# Configurando as colunas 'DENOM_CIA' e 'CNPJ_CIA' como índices nos DataFrames
deprec_amortz.set_index(['DENOM_CIA', 'CNPJ_CIA'], inplace=True)
LAIR.set_index(['DENOM_CIA', 'CNPJ_CIA'], inplace=True)


# Substituindo valores NaN por 0
deprec_amortz.fillna(0, inplace=True)
LAIR.fillna(0, inplace=True)

# Calculando o EBTIDA (soma das colunas deprec_amortz, juros e LAIR)
EBTIDA = LAIR.sub(deprec_amortz, fill_value=0)

# Resetando o índice para tornar 'DENOM_CIA' e 'CNPJ_CIA' em colunas novamente
EBTIDA.reset_index(inplace=True)

# %% [markdown]
# Lucro Bruto

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_DRE_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_dre = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar a Receita Operacional Líquida
    var_DRE = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '3.03')]

    # Selecionando colunas desejadas
    var_DRE = var_DRE[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_DRE['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_dre = pd.concat([resultados_dre, var_DRE], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_dre = resultados_dre.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
resultado_bruto = resultados_agregados_dre.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
resultado_bruto.name = None

# Selecionando as colunas desejadas
lucro_bruto = resultado_bruto[['DENOM_CIA', 'CNPJ_CIA'] + list(range(2012, 2023))]



# %% [markdown]
# Custo dos Produtos Vendidos (CPV)

# %%
# Configurando as colunas 'DENOM_CIA' e 'CNPJ_CIA' como índices nos DataFrames
vendas_liquidas.set_index(['DENOM_CIA', 'CNPJ_CIA'], inplace=True)
lucro_bruto.set_index(['DENOM_CIA', 'CNPJ_CIA'], inplace=True)


# Substituindo valores NaN por 0
vendas_liquidas.fillna(0, inplace=True)
lucro_bruto.fillna(0, inplace=True)

# Calculando o EBTIDA (soma das colunas deprec_amortz, juros e LAIR)
custo_prod_vendidos = vendas_liquidas.sub(lucro_bruto, fill_value=0)

# Resetando o índice para tornar 'DENOM_CIA' e 'CNPJ_CIA' em colunas novamente
custo_prod_vendidos.reset_index(inplace=True)

# %% [markdown]
# Estoques

# %%
# Definindo relatório a ser aberto
relatorio = "dfp_cia_aberta_BPA_ind"  # inserir apenas o nome do arquivo, sem o ano

# Lista de anos desejados
anos = list(range(2012, 2023))  # de 2012 a 2022

# DataFrame vazio para armazenar os resultados
resultados_bpa = pd.DataFrame()

# Loop pelos anos
for ano in anos:
    # Definindo o caminho do arquivo usando formatação de string
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'{relatorio}_{ano}.csv')

    # Lendo o arquivo CSV
    df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

    # Filtrando os dados para encontrar o Lucro Líquido do Período
    var_BPA = df[(df['ORDEM_EXERC'] == 'ÚLTIMO') & (df['CD_CONTA'] == '1.01.04')]

    # Selecionando colunas desejadas
    var_BPA = var_BPA[['DENOM_CIA', 'CNPJ_CIA', 'VL_CONTA']]

    # Adicionando a coluna 'Ano' com o valor correspondente
    var_BPA['Ano'] = ano

    # Adicionando os resultados ao DataFrame final
    resultados_bpa = pd.concat([resultados_bpa, var_BPA], ignore_index=True)

# Agregando os valores duplicados
resultados_agregados_bpa = resultados_bpa.groupby(['DENOM_CIA', 'CNPJ_CIA', 'Ano'])['VL_CONTA'].sum().reset_index()

# Criando um DataFrame pivotado
estoque = resultados_agregados_bpa.pivot_table(index=['DENOM_CIA', 'CNPJ_CIA'], columns='Ano', values='VL_CONTA', aggfunc='sum').reset_index()

# Renomeando as colunas para remover o nome da coluna do índice
estoque.name = None


estoque

# %% [markdown]
# Número de ações

# %% [markdown]
# Obtenção dos relatórios dos formulários de referencia (FRE). Há uma url para os relatórios dos ultimos 5 anos e uma url para os relatórios mais antigos, por isso a necessidade de separa em 2 urls os requerimentos.

# %%
# URL base dos relatórios até 2018
url_base_mais_5_anos = "https://dados.cvm.gov.br/dados/CIA_ABERTA/DOC/FRE/DADOS/"

# URL base dos relatórios a partir de 2019
url_base_ultimos_5_anos = "https://dados.cvm.gov.br/dataset/cia_aberta-doc-fre"

# Período
ano_atual = datetime.now().year
rel_anteriores = range(ano_atual - 11, ano_atual - 4)
rel_atuais = range(ano_atual - 4, ano_atual  + 1)

# Diretório de destino para download
diretorio_download = os.path.join(dir_arq_cvm_zip, f'Formulários de referencia (FRE)\Download')

# Diretório de destino para descompactação
diretorio_descompactacao = os.path.join(dir_arq_cvm_zip,f'Formulários de referencia (FRE)')

# Criando os diretórios se não existirem
os.makedirs(diretorio_download, exist_ok=True)
os.makedirs(diretorio_descompactacao, exist_ok=True)

# Iterando relatórios com mais de 5 anos
for ano in rel_anteriores:
    # Construindo o caminho do arquivo local para download
    caminho_arquivo_zip = os.path.join(diretorio_download, f"FRE_CIA_ABERTA_{ano}.zip")

    # Verificando se o arquivo já existe
    if os.path.exists(caminho_arquivo_zip):
        print(f"O arquivo para {ano} já foi baixado. Pulando para o próximo ano.")
    else:
        # Construindo a URL para o ano específico
        url_ano = f"{url_base_mais_5_anos}FRE_CIA_ABERTA_{ano}.zip"

        # Obtendo o conteúdo da página
        response = requests.get(url_ano, stream=True)

        # Verificando se a requisição foi bem-sucedida
        if response.status_code == 200:
            # Baixando o arquivo com barra de progresso
            with open(caminho_arquivo_zip, 'wb') as file, tqdm(
                desc=f"Fazendo download de {ano}",
                total=int(response.headers.get('content-length', 0)),
                unit='B',
                unit_scale=True,
                unit_divisor=1024,
            ) as bar:
                for data in response.iter_content(chunk_size=1024):
                    bar.update(len(data))
                    file.write(data)

            print(f"Download concluído para {ano}.")
        else:
            print(f"Falha ao obter o arquivo para {ano}.")

    # Construindo os caminhos dos arquivos
    diretorio_ano_descompactado = os.path.join(diretorio_descompactacao, str(ano))

    # Verificando se o diretório de descompactação já existe
    if os.path.exists(diretorio_ano_descompactado):
        print(f"Os arquivos para o ano {ano} já foram descompactados. Pulando para o próximo ano.")
    else:
        # Criando o diretório para o ano
        os.makedirs(diretorio_ano_descompactado, exist_ok=True)

        # Descompactando o arquivo ZIP no diretório do ano
        with zipfile.ZipFile(caminho_arquivo_zip, 'r') as zip_ref:
            zip_ref.extractall(diretorio_ano_descompactado)

        print(f"Descompactação concluída para {ano}.")

# Iterando sobre os anos a partir de 2019
for ano in rel_atuais:
    # Construindo o caminho do arquivo local para download
    caminho_arquivo_zip = os.path.join(diretorio_download, f"FRE_CIA_ABERTA_{ano}.zip")

    # Verificando se o arquivo já existe
    if os.path.exists(caminho_arquivo_zip):
        print(f"O arquivo para {ano} já foi baixado. Pulando para o próximo ano.")
    else:
        # Escolhendo a URL base correta com base no ano
        url_base = url_base_ultimos_5_anos if ano >= 2019 else url_base_mais_5_anos

        # Construindo a URL para o ano específico
        url_ano = f"{url_base}FRE_CIA_ABERTA_{ano}.zip"

        # Obtendo o conteúdo da página
        response = requests.get(url_ano, stream=True)

        # Verificando se a requisição foi bem-sucedida
        if response.status_code == 200:
            # Baixando o arquivo com barra de progresso
            with open(caminho_arquivo_zip, 'wb') as file, tqdm(
                desc=f"Fazendo download de {ano}",
                total=int(response.headers.get('content-length', 0)),
                unit='B',
                unit_scale=True,
                unit_divisor=1024,
            ) as bar:
                for data in response.iter_content(chunk_size=1024):
                    bar.update(len(data))
                    file.write(data)

            print(f"Download concluído para {ano}.")
        else:
            print(f"Falha ao obter o arquivo para {ano}.")

    # Construindo os caminhos dos arquivos
    diretorio_ano_descompactado = os.path.join(diretorio_descompactacao, str(ano))

    # Verificando se o diretório de descompactação já existe
    if os.path.exists(diretorio_ano_descompactado):
        print(f"Os arquivos para o ano {ano} já foram descompactados. Pulando para o próximo ano.")
    else:
        # Criando o diretório para o ano
        os.makedirs(diretorio_ano_descompactado, exist_ok=True)

        # Descompactando o arquivo ZIP no diretório do ano
        with zipfile.ZipFile(caminho_arquivo_zip, 'r') as zip_ref:
            zip_ref.extractall(diretorio_ano_descompactado)

        print(f"Descompactação concluída para {ano}.")

print("Downloads e descompactações concluídos.")


# %%
# Lista com todos os arquivos no diretório disponíveis para visualização
diretorio = os.path.join(dir_arq_cvm_zip, f'Formulários de referencia (FRE)/2013')
arquivos_disponiveis = os.listdir(diretorio)

# Exibe a lista de arquivos disponíveis
for arquivo in arquivos_disponiveis:
    print(arquivo)

# %%
# Visualizando relatório a ser aberto
relatorio = "fre_cia_aberta_endividamento"  # Inserir apenas o nome do arquivo, sem o ano
ano = 2019

# Definindo o caminho do arquivo usando formatação de string
caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'Formulários de referencia (FRE)/{ano}/{relatorio}_{ano}.csv')

df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')
df

# %% [markdown]
# Número de ações emitidas

# %%
# Definindo relatório a ser aberto
anos = range(2012, 2023)

# Criando um DataFrame vazio para armazenar os dados
df_numero_acoes = pd.DataFrame()

# Iterando pelos anos
for ano in anos:
    # Construindo o caminho do arquivo
    caminho_arquivo = os.path.join(dir_arq_cvm_zip, f'Formulários de referencia (FRE)/{ano}/fre_cia_aberta_capital_social_{ano}.csv')
    
    # Verificando se o arquivo existe antes de tentar lê-lo
    if os.path.exists(caminho_arquivo):
        # Lendo o arquivo CSV
        df = pd.read_csv(caminho_arquivo, encoding='cp1252', sep=';')

        # Filtrando o DataFrame para as colunas desejadas e o Tipo_Capital correto
        df_filtrado = df[df['Tipo_Capital'] == 'Capital Emitido'][['CNPJ_Companhia', 'Quantidade_Total_Acoes']]

        # Agrupando por CNPJ e somando as quantidades de ações para cada CNPJ
        df_filtrado = df_filtrado.groupby('CNPJ_Companhia')['Quantidade_Total_Acoes'].sum().reset_index()

        # Renomeando a coluna Quantidade_Total_Acoes para o ano correspondente
        df_filtrado = df_filtrado.rename(columns={'Quantidade_Total_Acoes': str(ano)})

        # Adicionando os dados ao DataFrame final
        if df_numero_acoes.empty:
            df_numero_acoes = df_filtrado
        else:
            df_numero_acoes = pd.merge(df_numero_acoes, df_filtrado, on='CNPJ_Companhia', how='outer')

df_numero_acoes
caminho_salvar = os.path.join(dir_out_n_acoes, f'numero_acoes_emitidas.csv')
df_numero_acoes.to_csv(caminho_salvar)

# %% [markdown]
# ==================================================


